"""
agents/word_document_generator.py
──────────────────────────────────
Agent 6 – Word Document Generator  (IEEE Two-Column Format)

Generates a properly formatted IEEE-style two-column research paper:
  • Section 1  (single-column): Title, Authors, Affiliation
  • Section 2  (two-column):    Abstract, Keywords, all body sections, References

Uses docx-js (Node.js) under the hood for reliable multi-column support.
Falls back to python-docx single-column if Node is unavailable.
"""

import os, sys, json, subprocess, tempfile
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from config.config import OUTPUT_DIR, OUTPUT_FILENAME

# ── Node.js / docx-js IEEE renderer ──────────────────────────────────────────
_JS_TEMPLATE = r"""
const {
  Document, Packer, Paragraph, TextRun,
  AlignmentType, SectionType, BorderStyle,
  Header, Footer, PageNumber, Column
} = require('docx');
const fs = require('fs');

const d = JSON.parse(fs.readFileSync(__JSONPATH__, 'utf8'));
const ROMAN = ['I','II','III','IV','V','VI','VII','VIII','IX','X','XI','XII','XIII','XIV'];

function bodyParas(text) {
  return text.split(/\n\n+/).filter(c => c.trim()).map(chunk =>
    new Paragraph({
      alignment: AlignmentType.JUSTIFIED,
      spacing: { after: 80, line: 240 },
      indent: { firstLine: 360 },
      children: [new TextRun({ text: chunk.replace(/\n/g,' '), font:'Times New Roman', size:20 })]
    })
  );
}

function sectionHead(num, title) {
  return new Paragraph({
    alignment: AlignmentType.CENTER,
    spacing: { before: 160, after: 80 },
    children: [new TextRun({ text: `${num}.  ${title.toUpperCase()}`,
      font:'Times New Roman', size:20, bold:true })]
  });
}

function refLine(text) {
  return new Paragraph({
    alignment: AlignmentType.JUSTIFIED,
    spacing: { after: 60 },
    indent: { left:360, hanging:360 },
    children: [new TextRun({ text, font:'Times New Roman', size:18 })]
  });
}

// ── Section 1: full-width title block ────────────────────────────────────────
const titleSection = {
  properties: {
    page: {
      size: { width:12240, height:15840 },
      margin: { top:1080, bottom:1080, left:1080, right:1080 },
    },
    type: SectionType.CONTINUOUS,
    column: { count:1 },
  },
  headers: {
    default: new Header({ children:[
      new Paragraph({
        alignment: AlignmentType.CENTER,
        border: { bottom:{ style:BorderStyle.SINGLE, size:6, color:'000000', space:1 }},
        children:[new TextRun({ text: d.conference_line || 'International Journal of Engineering Research',
          font:'Times New Roman', size:16, italics:true })]
      })
    ]})
  },
  footers: {
    default: new Footer({ children:[
      new Paragraph({
        alignment: AlignmentType.CENTER,
        children:[
          new TextRun({ text:'Page ', font:'Times New Roman', size:18 }),
          new TextRun({ children:[PageNumber.CURRENT], font:'Times New Roman', size:18 }),
          new TextRun({ text:' of ', font:'Times New Roman', size:18 }),
          new TextRun({ children:[PageNumber.TOTAL_PAGES], font:'Times New Roman', size:18 }),
        ]
      })
    ]})
  },
  children: [
    new Paragraph({
      alignment: AlignmentType.CENTER,
      spacing: { before:200, after:120 },
      children:[new TextRun({ text:d.title, font:'Times New Roman', size:36, bold:true })]
    }),
    new Paragraph({
      alignment: AlignmentType.CENTER,
      spacing: { after:60 },
      children:[new TextRun({ text:d.authors, font:'Times New Roman', size:20, bold:true })]
    }),
    new Paragraph({
      alignment: AlignmentType.CENTER,
      spacing: { after:200 },
      children:[new TextRun({ text:d.affiliation, font:'Times New Roman', size:18, italics:true })]
    }),
  ]
};

// ── Section 2: two-column body ───────────────────────────────────────────────
const body = [];

// Abstract
body.push(new Paragraph({
  alignment: AlignmentType.CENTER,
  spacing:{ before:120, after:80 },
  children:[new TextRun({ text:'Abstract', font:'Times New Roman', size:20, bold:true, italics:true })]
}));
body.push(new Paragraph({
  alignment: AlignmentType.JUSTIFIED,
  spacing:{ after:100 },
  children:[
    new TextRun({ text:'Abstract\u2014', font:'Times New Roman', size:18, bold:true, italics:true }),
    new TextRun({ text:d.abstract_text, font:'Times New Roman', size:18 }),
  ]
}));
body.push(new Paragraph({
  alignment: AlignmentType.JUSTIFIED,
  spacing:{ after:140 },
  children:[
    new TextRun({ text:'Index Terms\u2014', font:'Times New Roman', size:18, bold:true, italics:true }),
    new TextRun({ text:d.keywords, font:'Times New Roman', size:18, italics:true }),
  ]
}));
body.push(new Paragraph({
  spacing:{ after:100 },
  border:{ bottom:{ style:BorderStyle.SINGLE, size:6, color:'000000', space:2 }},
  children:[]
}));

// Body sections
let idx = 0;
for (const sec of d.sections) {
  const isRef = sec.heading.toLowerCase().includes('reference');
  body.push(sectionHead(ROMAN[idx] || String(idx+1), sec.heading));
  idx++;
  if (isRef) {
    sec.text.split('\n').filter(l=>l.trim()).forEach(l => body.push(refLine(l)));
  } else {
    bodyParas(sec.text).forEach(p => body.push(p));
  }
}

const bodySection = {
  properties: {
    page: {
      size:{ width:12240, height:15840 },
      margin:{ top:1080, bottom:1080, left:1080, right:1080 },
    },
    type: SectionType.CONTINUOUS,
    column: { count:2, space:720, equalWidth:true, separate:false },
  },
  children: body,
};

const doc = new Document({ sections:[titleSection, bodySection] });
Packer.toBuffer(doc).then(buf => {
  fs.writeFileSync(__OUTPATH__, buf);
  console.log('OK');
}).catch(e => { console.error(e.message); process.exit(1); });
"""


def _run_node_renderer(payload: dict, out_path: str) -> bool:
    """Write JSON, inject paths into JS template, run node. Returns True on success."""
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json',
                                         delete=False, encoding='utf-8') as jf:
            json.dump(payload, jf, ensure_ascii=False)
            json_path = jf.name

        with tempfile.NamedTemporaryFile(mode='w', suffix='.js',
                                         delete=False, encoding='utf-8') as jsf:
            js = _JS_TEMPLATE.replace('__JSONPATH__', repr(json_path)) \
                             .replace('__OUTPATH__',  repr(out_path))
            jsf.write(js)
            js_path = jsf.name

        result = subprocess.run(['node', js_path], capture_output=True, text=True, timeout=60)
        if result.returncode == 0 and os.path.isfile(out_path):
            return True
        print(f"[WordGen] Node error: {result.stderr[:300]}")
        return False
    except Exception as e:
        print(f"[WordGen] Node renderer failed: {e}")
        return False
    finally:
        for p in [json_path, js_path]:
            try: os.unlink(p)
            except: pass


def _fallback_python_docx(payload: dict, out_path: str):
    """Simple python-docx fallback (single-column) if Node is unavailable."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.75)

    def _p(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        return p

    _p(payload['title'], bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(payload['authors'], bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(payload['affiliation'], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(f"Abstract— {payload['abstract_text']}", size=10)
    _p(f"Index Terms— {payload['keywords']}", size=10)
    ROMAN = ['I','II','III','IV','V','VI','VII','VIII','IX','X','XI','XII']
    for i, sec in enumerate(payload['sections']):
        _p(f"{ROMAN[i] if i < len(ROMAN) else i+1}.  {sec['heading'].upper()}",
           bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        for chunk in sec['text'].split('\n\n'):
            if chunk.strip():
                _p(chunk.strip(), size=10)
    doc.save(out_path)


def generate_word_document(
    inputs: dict,
    template: dict,
    section_texts: dict,
    references: list,
    formatted_refs: list,
    authors: str = "Research Author, Co-Author",
) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.abspath(os.path.join(OUTPUT_DIR, OUTPUT_FILENAME))

    # Build sections list (exclude Abstract/Keywords — handled separately)
    skip = {'Abstract', 'Keywords'}
    sections_data = []
    for sec in template['sections']:
        if sec in skip:
            continue
        if sec.lower() == 'references':
            sections_data.append({
                'heading': 'References',
                'text': '\n'.join(formatted_refs)
            })
        else:
            sections_data.append({
                'heading': sec,
                'text': section_texts.get(sec, '')
            })

    payload = {
        'title':           inputs['title'],
        'authors':         authors,
        'affiliation':     'Department of Electronics and Communication Engineering, '
                           'KIT – Kalaignarkarunanidhi Institute of Technology, Coimbatore',
        'conference_line': 'International Journal of Engineering Research & Technology (IJERT)',
        'abstract_text':   section_texts.get('Abstract', ''),
        'keywords':        section_texts.get('Keywords', ''),
        'sections':        sections_data,
    }

    print('[WordDocGenerator] Rendering IEEE two-column format …')
    success = _run_node_renderer(payload, out_path)

    if not success:
        print('[WordDocGenerator] Falling back to python-docx single-column …')
        _fallback_python_docx(payload, out_path)

    print(f'[WordDocGenerator] Saved → {out_path}')
    return out_path